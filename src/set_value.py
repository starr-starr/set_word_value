import docx

from utils.utils import load_json


def set_value(path):
    file = docx.Document(path)
    paragraphs = file.paragraphs[1]
    paragraphs.clear()

    data = load_json()
    for key, value in data.items():
        key_run = paragraphs.add_run()
        value_run = paragraphs.add_run()
        key_run.text = key
        value_run.text = value
        value_run.font.underline = True

    file.save(path)
